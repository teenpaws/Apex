"""
Unit tests for DocumentExtractor service.

Tests cover:
  - DOCX text extraction from fixture bytes
  - Empty bytes returns empty string (not an error)
  - Unsupported file type raises ValueError
  - Whitespace normalisation (multiple spaces collapsed)
"""
from __future__ import annotations

import io
import pytest

_MODULE = "app.services.document_extractor"


def _make_minimal_docx() -> bytes:
    """Minimal DOCX with the text 'Hello DOCX' using python-docx."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Hello DOCX")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocumentExtractor:

    @pytest.fixture(autouse=True)
    def _import(self):
        pytest.importorskip(_MODULE)

    def test_extract_docx_contains_text(self):
        from app.services.document_extractor import DocumentExtractor
        docx_bytes = _make_minimal_docx()
        text = DocumentExtractor.extract(docx_bytes, "DOCX")
        assert "Hello DOCX" in text

    def test_extract_empty_pdf_returns_empty_string(self):
        from app.services.document_extractor import DocumentExtractor
        # Empty bytes — pdfplumber raises, we return ""
        text = DocumentExtractor.extract(b"", "PDF")
        assert text == ""

    def test_extract_unsupported_type_raises(self):
        from app.services.document_extractor import DocumentExtractor
        with pytest.raises(ValueError, match="Unsupported file type"):
            DocumentExtractor.extract(b"data", "TXT")

    def test_extract_docx_normalises_whitespace(self):
        from app.services.document_extractor import DocumentExtractor
        from docx import Document
        doc = Document()
        doc.add_paragraph("  Hello   World  ")
        buf = io.BytesIO()
        doc.save(buf)
        text = DocumentExtractor.extract(buf.getvalue(), "DOCX")
        # Multiple consecutive spaces should be collapsed
        assert "  " not in text

    def test_extract_pdf_returns_string(self):
        from app.services.document_extractor import DocumentExtractor
        # Valid call with empty PDF returns empty string (not error)
        result = DocumentExtractor.extract(b"", "PDF")
        assert isinstance(result, str)
